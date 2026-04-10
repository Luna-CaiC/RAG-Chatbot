"""
src/ingest.py — Document Processing Pipeline
=============================================
Handles file upload → text extraction → chunking → embedding → vector store.
Supports: PDF (text + image-based OCR), Word (.docx), Images (.jpg/.png/etc.)
Uses a local HuggingFace embedding model (no API rate limits).
"""

import os
import io
import time
import uuid
import logging
import tempfile
import fitz
import base64

import docx
from PIL import Image

from dotenv import load_dotenv
from langchain_community.document_loaders import PyMuPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import HumanMessage

load_dotenv(override=True)

# ── Constants ────────────────────────────────────────────────────────
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200
EMBEDDING_MODEL = "all-MiniLM-L6-v2"   # runs locally, no API calls
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.0-flash")

# File type groups
PDF_TYPES   = {".pdf"}
WORD_TYPES  = {".docx"}
IMAGE_TYPES = {".jpg", ".jpeg", ".png", ".webp", ".gif"}

logger = logging.getLogger(__name__)


# ── Helpers ──────────────────────────────────────────────────────────

def _call_gemini_vision(llm: ChatGoogleGenerativeAI, prompt: str, images_b64: list) -> str | None:
    """
    Send one or more base64 images to Gemini Vision with retry on rate limit.
    Returns the response text, or None if all retries fail.
    """
    image_parts = [
        {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img}"}}
        for img in images_b64
    ]
    msg = HumanMessage(content=[{"type": "text", "text": prompt}] + image_parts)

    for attempt in range(3):
        try:
            return llm.invoke([msg]).content
        except Exception as e:
            if "429" in str(e) and attempt < 2:
                wait = 35 * (attempt + 1)
                logger.warning("Vision API rate limited — retrying in %ds...", wait)
                time.sleep(wait)
            else:
                logger.warning("Vision API call failed: %s", e)
                break
    return None


def _extract_pdf_pages(file_bytes: bytes, filename: str, llm: ChatGoogleGenerativeAI) -> list[Document]:
    """Extract text from a PDF. Uses Gemini Vision OCR for image-based PDFs."""
    pages = []
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        loader = PyMuPDFLoader(tmp_path)
        pages = loader.load()

        # Detect image-based PDF (avg < 50 chars/page)
        total_chars = sum(len(p.page_content.strip()) for p in pages)
        avg_chars = total_chars / max(1, len(pages))

        if avg_chars < 50:
            logger.info("'%s' is image-based (avg %d chars/page). Running OCR...", filename, avg_chars)
            pdf_doc = fitz.open(tmp_path)
            ocr_images, ocr_indices = [], []

            for i, page in enumerate(pages):
                if len(page.page_content.strip()) <= 50:
                    ocr_indices.append(i)
                    pix = pdf_doc[i].get_pixmap(dpi=100)
                    ocr_images.append(base64.b64encode(pix.tobytes("png")).decode("utf-8"))

            if ocr_images:
                batch_size = 15
                total_batches = (len(ocr_images) + batch_size - 1) // batch_size
                for batch_num, start in enumerate(range(0, len(ocr_images), batch_size)):
                    batch_imgs = ocr_images[start:start + batch_size]
                    batch_idx  = ocr_indices[start:start + batch_size]
                    logger.info("OCR batch %d/%d for '%s'...", batch_num + 1, total_batches, filename)

                    ocr_prompt = (
                        "Extract all text from these slide images verbatim. "
                        "Prefix each slide's text with 'Slide N:' where N is its order. "
                        "Summarize charts and diagrams as text."
                    )
                    ocr_text = _call_gemini_vision(llm, ocr_prompt, batch_imgs)
                    if ocr_text:
                        parts = [s.strip() for s in ocr_text.split("Slide ") if s.strip()]
                        for j, idx in enumerate(batch_idx):
                            pages[idx].page_content = f"Slide {parts[j]}" if j < len(parts) else ocr_text
                    else:
                        logger.warning("OCR failed for batch %d of '%s'.", batch_num + 1, filename)

            pdf_doc.close()
        else:
            logger.info("'%s' is a text PDF. Skipping OCR.", filename)
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)

    return pages


def _extract_word_pages(file_bytes: bytes, filename: str) -> list[Document]:
    """
    Extract text from a Word (.docx) file.
    Method 1: python-docx paragraphs + tables (standard)
    Method 2: raw XML <w:t> tag extraction (fallback for text boxes, etc.)
    """
    import zipfile
    import re as _re

    text_parts = []

    # ── Method 1: python-docx ────────────────────────────────────────
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)
        logger.info("python-docx extracted %d text blocks from '%s'.", len(text_parts), filename)
    except Exception as e:
        logger.warning("python-docx failed for '%s': %s — trying XML fallback.", filename, e)

    # ── Method 2: raw XML fallback (handles text boxes, shapes, etc.) ─
    if not text_parts:
        logger.info("Trying XML fallback for '%s'...", filename)
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                xml_files = [n for n in z.namelist()
                             if n.startswith("word/") and n.endswith(".xml")]
                for xml_name in xml_files:
                    xml_content = z.read(xml_name).decode("utf-8", errors="ignore")
                    # Extract text between <w:t> tags
                    tags = _re.findall(r'<w:t[^>]*>([^<]+)</w:t>', xml_content)
                    text_parts.extend(t.strip() for t in tags if t.strip())
            logger.info("XML fallback extracted %d text fragments from '%s'.",
                        len(text_parts), filename)
        except Exception as e:
            logger.warning("XML fallback also failed for '%s': %s", filename, e)

    combined = "\n".join(text_parts).strip()
    if not combined:
        logger.warning("No text extracted from Word file '%s'.", filename)
        return []

    return [Document(page_content=combined, metadata={"source": filename, "page": 0})]


def _extract_image_pages(file_bytes: bytes, filename: str, llm: ChatGoogleGenerativeAI) -> list[Document]:
    """
    Extract text from a standalone image file using Gemini Vision OCR.
    """
    # Convert image to PNG bytes for consistent base64 encoding
    img = Image.open(io.BytesIO(file_bytes)).convert("RGB")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    img_b64 = base64.b64encode(buf.getvalue()).decode("utf-8")

    logger.info("Running Vision OCR on image '%s'...", filename)
    ocr_prompt = (
        "Extract all text visible in this image verbatim. "
        "If the image contains a diagram or chart, describe it in detail. "
        "If there is no text, describe the image content."
    )
    ocr_text = _call_gemini_vision(llm, ocr_prompt, [img_b64])

    if not ocr_text:
        logger.warning("Vision OCR failed for image '%s'.", filename)
        return []

    return [Document(page_content=ocr_text, metadata={"source": filename, "page": 0})]


# ── Main entry point ─────────────────────────────────────────────────

def process_documents(uploaded_files: list) -> Chroma:
    """
    Process multiple uploaded files (PDF / Word / Image) into a combined vector store.

    Args:
        uploaded_files: List of Streamlit UploadedFile objects.

    Returns:
        Chroma: A freshly created, session-isolated vector store.
    """
    if not uploaded_files:
        raise ValueError("No files were uploaded.")

    llm = ChatGoogleGenerativeAI(model=GEMINI_MODEL, temperature=0.1)
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP
    )
    all_chunks = []

    for uploaded_file in uploaded_files:
        file_bytes = uploaded_file.read()
        if not file_bytes:
            logger.warning("Skipping empty file: %s", uploaded_file.name)
            continue

        name = uploaded_file.name
        ext  = os.path.splitext(name)[1].lower()

        try:
            # ── Route by file type ────────────────────────────────────
            if ext in PDF_TYPES:
                pages = _extract_pdf_pages(file_bytes, name, llm)
            elif ext in WORD_TYPES:
                pages = _extract_word_pages(file_bytes, name)
            elif ext in IMAGE_TYPES:
                pages = _extract_image_pages(file_bytes, name, llm)
            else:
                logger.warning("Unsupported file type '%s' — skipping.", name)
                continue
            # ─────────────────────────────────────────────────────────

            if not pages:
                logger.warning("No content extracted from '%s'.", name)
                continue

            # Set source metadata
            for page in pages:
                page.metadata["source"] = name

            # Split into chunks
            chunks = text_splitter.split_documents(pages)

            # ── Inject source label into EVERY chunk (post-split) ─────
            # Critical: without this, only the first chunk of each page
            # carries the filename — later chunks are anonymous.
            source_label = f"[Source Document: {name}]"
            valid_chunks = []
            for chunk in chunks:
                content = chunk.page_content.strip()
                if len(content) < 5:
                    continue  # discard truly empty chunks only
                if not content.startswith("[Source Document:"):
                    chunk.page_content = f"{source_label}\n{content}"
                valid_chunks.append(chunk)
            # ─────────────────────────────────────────────────────────

            all_chunks.extend(valid_chunks)
            logger.info("Loaded %d chunks from '%s' (%s).", len(valid_chunks), name, ext)

        except Exception as exc:
            logger.error("Failed to process '%s': %s", name, exc)
            raise RuntimeError(f"Failed to process '{name}': {exc}") from exc

    if not all_chunks:
        raise RuntimeError("No text could be extracted from any of the uploaded files. "
                           "The file may be empty, password-protected, or use an unsupported format.")

    embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)

    # ── Unique collection per session — prevents cross-session contamination ──
    collection_name = f"rag_session_{uuid.uuid4().hex[:12]}"
    vector_store = Chroma.from_documents(
        documents=all_chunks,
        embedding=embeddings,
        collection_name=collection_name,
    )

    logger.info("Vector store '%s' created: %d vectors from %d file(s).",
                collection_name, len(all_chunks), len(uploaded_files))
    return vector_store
